#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(1.8);  sec.bottom_margin = Cm(1.8)

C_WHITE = RGBColor(0xFF,0xFF,0xFF); C_DARK = RGBColor(0x1E,0x29,0x3B)
C_BLUE  = RGBColor(0x00,0x72,0xA8); C_GRAY = RGBColor(0x64,0x74,0x8B)
C_GREEN = RGBColor(0x05,0x96,0x69); C_ORANGE = RGBColor(0xEA,0x58,0x0C)
C_RED   = RGBColor(0xDC,0x26,0x26)
GANTT_BLUE="1A5276"; GANTT_MS="C0392B"; GANTT_EMPTY="F8FAFC"
HEADER_BG="1E293B"; SUBHDR_BG="2E4057"

MONTHS=[
    ("Ay 1","Sub 2026"),("Ay 2","Mar 2026"),("Ay 3","Nis 2026"),
    ("Ay 4","May 2026"),("Ay 5","Haz 2026"),("Ay 6","Tem 2026"),
    ("Ay 7","Ago 2026"),("Ay 8","Eyl 2026"),("Ay 9","Eki 2026"),
    ("Ay 10","Kas 2026"),("Ay 11","Ara 2026"),("Ay 12","Oca 2027"),
    ("Ay 13","Sub 2027"),("Ay 14","Mar 2027"),("Ay 15","Nis 2027"),
    ("Ay 16","May 2027"),("Ay 17","Haz 2027"),("Ay 18","Tem 2027"),
]

IP=[
  {"kod":"IP1","ad":"Temel Platform Altyapisi","start":0,"end":1,"ms":[1,3],"renk":"154360",
   "ozet":"Projenin teknik temelini olusturan bu is paketi; Docker Compose ortamini, FastAPI/TimescaleDB/Redis mimarisini, JWT kimlik dogrulama sistemini, temel sunucu CRUD API'sini ve React+Vite frontend iskeletini kapsar.",
   "ciktilar":["Docker Compose stack (TimescaleDB, Redis, Prometheus, Pushgateway)","FastAPI backend + SQLAlchemy 2.x + Alembic migration","JWT kimlik dogrulama ve CORS yapilandirmasi","Server CRUD API (servers.py) + SSH Manager (Paramiko)","React 18 + Vite + TypeScript frontend iskeleti","Layout, routing, Dashboard ve LiveMetrics sayfasi","Background task engine (5 asyncio gorevi)","Settings API (AppSettings, credentials, RAG toggle)"],
   "aylik":{0:["Docker Compose stack kurulumu","FastAPI + Alembic iskelet","TimescaleDB sema tasarimi","JWT auth modulu","React+Vite+TS frontend kurulumu"],1:["Server CRUD API","SSH Manager (Paramiko baglanti havuzu)","Node Exporter kurulum servisi (SCP/HTTP/Base64)","Background task engine","Prometheus entegrasyonu + metric sync","Dashboard + Servers + LiveMetrics sayfasi","Settings API"]},
   "kta":"docker compose up ile 5 dk icinde tum servisler ayaga kalkar; sunucu CRUD calisir, LiveMetrics Prometheus verisini gosteriyor."},
  {"kod":"IP2","ad":"Envanter ve Kesif","start":1,"end":3,"ms":[2,3],"renk":"1A5276",
   "ozet":"Sunucu ve sanal makine envanterinin otomatik olarak kesfedilmesi ve senkronize edilmesi. VMware vCenter ve oVirt hipervizorlerinden VM listesi cekilir; linux_info_collector ile donanim/OS detaylari kayit altina alinir.",
   "ciktilar":["linux_info_collector.py: SSH ile CPU/RAM/Disk/Network/OS","VMware vCenter client (vcenter_client.py)","oVirt client (ovirt_client.py)","Envanter senkronizasyon servisi (inventory_sync_service.py)","Hypervisors sayfasi (VM listesi, durum)","SSH key deployer (ssh_key_deployer.py)","Periyodik envanter sync background task"],
   "aylik":{1:["linux_info_collector tamamlama","Prometheus hedef yonetimi (prometheus_target_manager.py)"],2:["VMware vCenter API entegrasyonu","oVirt client gelistirme","SSH key deployer"],3:["inventory_sync_service.py","Hypervisors UI sayfasi","Entegrasyon testleri + milestone"]},
   "kta":"VMware/oVirt envanteri tam senkronize; Hypervisors sayfasinda VM'ler gozukuyor."},
  {"kod":"IP3","ad":"Log Toplama & Analiz","start":2,"end":4,"ms":[4],"renk":"1F618D",
   "ozet":"Sunuculardan periyodik SSH log toplama, merkezi depolama ve temel analiz. syslog ve uygulama loglari toplanir; Log sayfasi uzerinden goruntulenir ve filtrelenir.",
   "ciktilar":["log_collector.py: 15 dakikalık SSH log toplama","Log veritabani modeli (TimescaleDB)","Log filtreleme ve arama API","Log.tsx sayfasi (App.tsx routing ile)","Log anomali on-isleme altyapisi"],
   "aylik":{2:["log_collector.py SSH log toplama","Log veritabani modeli"],3:["Log API endpoint'leri","Log.tsx sayfasi + App.tsx routing duzeltmesi"],4:["Log arama ve filtreleme","Log anomali on-isleme","Milestone: Log modulu tamamlandi"]},
   "kta":"15 dakikada bir log toplanir; Log sayfasinda arama/filtreleme calisir."},
  {"kod":"IP4","ad":"Anomali Tespiti","start":3,"end":7,"ms":[4,7],"renk":"21618C",
   "ozet":"Prometheus metrikleri ve loglar uzerinde istatistiksel anomali tespiti. Otomatik uyari uretimi, AIOps event olusturma ve AnomalyDetection sayfasi gelistirme.",
   "ciktilar":["anomaly_detector.py: Prometheus 5 dk tarama","log_anomaly_detector.py","AIOps events API genisletme","incident_auto.py: otomatik incident","AnomalyDetection sayfasi","E-posta/webhook bildirim altyapisi"],
   "aylik":{3:["Prometheus anomali tarama altyapisi","Anomali model parametreleri"],4:["Log anomali detector entegrasyonu","Milestone: ilk anomali algilandi"],5:["AIOps event motoru genisletme","incident_auto.py otomatik incident"],6:["Bildirim sistemi (email/webhook)","Korelasyon motoru"],7:["AnomalyDetection UI gelistirme","Performans tuning","Milestone: anomali modulu tamamlandi"]},
   "kta":"CPU/RAM spike anomalileri 5 dk icinde tespit edilir; otomatik incident olusur."},
  {"kod":"IP5","ad":"Kapasite Planlama","start":5,"end":8,"ms":[7],"renk":"2874A6",
   "ozet":"Tarihi metrik verileri analiz ederek kaynak kullanim tahminleri uretmek. Kapasite raporu, trend analizi ve 30/60/90 gunluk buyume projeksiyonlarini kapsar.",
   "ciktilar":["Kapasite analiz servisi (TimescaleDB sorgulari)","Trend analizi: CPU/RAM/Disk projeksiyon","Kapasite raporu API ve UI sayfasi","Esik deger uyarilari","30/60/90 gun kaynak buyume tahmini"],
   "aylik":{5:["Kapasite veri modeli","Trend analiz algoritmalari (moving average)"],6:["Kapasite API endpoint'leri","TimescaleDB sorgu optimizasyonu"],7:["Kapasite UI + grafik gorsellestirme","Milestone: kapasite raporu aktif"],8:["Esik deger uyarilari","30/60/90 gun projeksiyon motoru"]},
   "kta":"30/60/90 gunluk kapasite projeksiyonu UI'da gozukuyor; esik uyarilari aktif."},
  {"kod":"IP6","ad":"Servis Haritalama","start":7,"end":11,"ms":[9,11],"renk":"2E86C1",
   "ozet":"Sunucular arasi bagimliliklarin ve servis iliskilerinin otomatik kesfedilerek haritaya dokulur. Topoloji goruntuleme, bagimlilik ve etki analizi altyapisi saglar.",
   "ciktilar":["Servis kesif motoru (network scan + SSH)","Bagimlilik veritabani modeli","Servis haritasi API","Topoloji goruntuleme UI","Etki analizi: hangi servis ne etkiler","Servis saglik skoru"],
   "aylik":{7:["Servis kesif altyapisi (port tarama + proses listesi)"],8:["Bagimlilik veri modeli + API"],9:["Topoloji goruntuleme UI","Milestone: servis haritasi v1"],10:["Etki analizi motoru","Servis saglik skoru"],11:["Anomali + servis haritasi entegrasyonu","Milestone: etki analizi aktif"]},
   "kta":"Sunucular arasi bagimliliklar haritada gozukur; ariza etki analizi calisir."},
  {"kod":"IP7","ad":"AI Chat & Otonom Operasyon","start":10,"end":15,"ms":[12,15],"renk":"1A6DA0",
   "ozet":"Ollama/bulut LLM destekli streaming AI sohbet, RAG (ChromaDB) ile zenginlestirilmis sorgu motoru, MCP Linux araclari ve otonom remediation is akislari. Ansible entegrasyonu ile chat uzerinden operasyon tetiklenebilir.",
   "ciktilar":["Ollama streaming chat (chat.py)","ChromaDB RAG servisi (runbook/incident/metric)","Multi-LLM destegi (Groq, OpenAI, Anthropic, OpenRouter)","MCP Linux Tools (mcp_client.py + McpTools sayfasi)","AI Agent otonom is akislari","Chat -> Ansible baglantisi","PDF parser + embedding servisi","Ansible/AWX UI sayfasi"],
   "aylik":{10:["Ollama chat entegrasyonu","Streaming Chat UI (Chat.tsx)"],11:["ChromaDB RAG kurulumu","Runbook/incident PDF ingest + embedding"],12:["Multi-LLM fallback (Groq/OpenAI/Anthropic)","Milestone: RAG calisir; chat context-aware"],13:["MCP Linux Tools servisi","McpTools sayfasi"],14:["AI Agent otonom is akislari","Ansible/AWX entegrasyonu (ansible.py)"],15:["Chat -> Ansible baglantisi","Milestone: otonom remediation aktif"]},
   "kta":"Chat'ten 'sunucu X disk temizle' dedikten sonra Ansible otomatik calisir."},
  {"kod":"IP8","ad":"Test & Urunlestirme","start":13,"end":17,"ms":[14,16,17],"renk":"154360",
   "ozet":"Urun kalite guvencesi, guvenlik sertlestirme, CI/CD boru hatti, kapsamli dokumantasyon ve GA (General Availability) v1.0 surum hazirligini kapsar.",
   "ciktilar":["E2E test suite (Playwright)","Unit test coverage > %70 (pytest)","CI/CD pipeline (GitHub Actions)","Guvenlik hardening (OWASP checklist)","OpenAPI dokumantasyonu","Deployment guide + README guncelleme","GA Release v1.0"],
   "aylik":{13:["E2E test altyapisi (Playwright)","Kritik akis testleri"],14:["Unit testler (pytest) + coverage raporu","Milestone: test coverage > %50"],15:["CI/CD pipeline (GitHub Actions)","Guvenlik hardening (OWASP)"],16:["Milestone: coverage > %70 & CI/CD aktif","OpenAPI dokumantasyonu"],17:["Deployment guide","Son hata duzeltmeleri","Milestone: GA Release v1.0"]},
   "kta":"%70+ test coverage; CI/CD aktif; deployment guide mevcut; v1.0 yayinda."},
]

def shd(cell, hex_str):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hex_str); tcPr.append(s)

def brd(cell, color="CBD5E1"):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcB=OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el=OxmlElement(f'w:{side}'); el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'4'); el.set(qn('w:space'),'0'); el.set(qn('w:color'),color); tcB.append(el)
    tcPr.append(tcB)

def cr(cell, text, bold=False, size=8, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p=cell.paragraphs[0]; p.alignment=align; r=p.add_run(text)
    r.bold=bold; r.font.size=Pt(size); r.font.name='Calibri'
    if color: r.font.color.rgb=color

def para(doc, text, bold=False, italic=False, color=None, size=10,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    p=doc.add_paragraph(); p.alignment=align
    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    r=p.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.name='Calibri'
    if color: r.font.color.rgb=color

def hd(doc, text, lvl=1, color=None):
    p=doc.add_heading(text,level=lvl); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name='Calibri'
        if color: r.font.color.rgb=color

def hline(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),'CBD5E1')
    pBdr.append(b); pPr.append(pBdr)

def pb(doc): doc.add_page_break()

# ── KAPAK ─────────────────────────────────────────────────────────────────────
for _ in range(4):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)

para(doc,"datatem AI",bold=True,size=38,color=C_BLUE,align=WD_ALIGN_PARAGRAPH.CENTER,sa=4)
para(doc,"AIOps & Server Management Platform",size=14,color=C_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER,sa=6)
para(doc,"18 AYLIK PROJE PLANI",bold=True,size=22,color=C_DARK,align=WD_ALIGN_PARAGRAPH.CENTER,sa=4)
para(doc,"Subat 2026  -  Temmuz 2027  |  8 Is Paketi  |  18 Milestone",size=12,color=C_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER,sa=30)
t=doc.add_table(rows=5,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
info=[("Baslangic","Subat 2026"),("Bitis","Temmuz 2027"),("Sure","18 Ay / 8 Is Paketi"),("Durum","IP1 - Devam Ediyor"),("Surum","v1.0  |  Mart 2026")]
for i,(k,v) in enumerate(info):
    c=t.rows[i].cells; c[0].width=Cm(5); c[1].width=Cm(5)
    shd(c[0],"1E293B"); shd(c[1],"F8FAFC")
    cr(c[0],k,bold=True,size=9,color=RGBColor(0x94,0xA3,0xB8),align=WD_ALIGN_PARAGRAPH.RIGHT)
    cr(c[1],v,bold=True,size=9,color=C_DARK)
for _ in range(5):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
para(doc,"Gizlilik: Sirket Ici  |  Versiyon 1.0  |  Mart 2026",size=8,italic=True,color=C_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)
pb(doc)

# ── ICINDEKILER ───────────────────────────────────────────────────────────────
hd(doc,"ICINDEKILER",lvl=1,color=C_BLUE); hline(doc)
toc=[("1.","Proje Ozeti & Teknoloji Stack"),("2.","Is Paketi Aciklamalari (IP1-IP8)"),("3.","18 Aylik Gantt Tablosu"),("4.","Aylik Aktivite Plani (Ay 1-18)"),("5.","Milestone Ozeti"),("6.","Risk & Bagimlilik Matrisi"),("7.","Kaynak & Efor Ozeti")]
for num,title in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
    r1=p.add_run(f"  {num}  "); r1.bold=True; r1.font.size=Pt(10); r1.font.name='Calibri'; r1.font.color.rgb=C_BLUE
    r2=p.add_run(title); r2.font.size=Pt(10); r2.font.name='Calibri'; r2.font.color.rgb=C_DARK
pb(doc)

# ── 1. PROJE OZETI ────────────────────────────────────────────────────────────
hd(doc,"1. Proje Ozeti & Teknoloji Stack",lvl=1,color=C_BLUE); hline(doc)
para(doc,"datatem AI; Linux sunuculari merkezi yonetmek, Prometheus canli metrikler izlemek, AI destekli sohbet (RAG) ve anomali tespiti ile olaylari otomatik analiz etmek, Ansible/AWX ile otomasyon calistirmak ve VMware/oVirt hipervizorlerle envanter senkronizasyonu yapmak uzere tasarlanmis AIOps platformudur. 8 is paketi 18 aya yayilmaktadir.",size=10,sa=8)
stack=[("Backend","FastAPI + SQLAlchemy 2.x + Alembic + Uvicorn"),("Veritabani","TimescaleDB (PostgreSQL 15) + Redis"),("Frontend","React 18 + Vite 5 + TypeScript + Tailwind + TanStack Query + Recharts"),("Izleme","Prometheus + Pushgateway + Node Exporter"),("AI / RAG","Ollama (yerel) + ChromaDB + pypdf + Groq/OpenAI/Anthropic (yedek)"),("Otomasyon","Ansible + AWX + Paramiko SSH"),("Hipervizor","VMware vCenter + oVirt"),("Container","Docker Compose (tek komutla ayaga kalkan ortam)")]
t=doc.add_table(rows=len(stack),cols=2); t.style='Table Grid'
for i,(k,v) in enumerate(stack):
    c=t.rows[i].cells; c[0].width=Cm(3.8); c[1].width=Cm(18)
    shd(c[0],"1E293B"); shd(c[1],"F0F9FF" if i%2==0 else "FFFFFF")
    brd(c[0],"334155"); brd(c[1],"E0F2FE")
    cr(c[0],k,bold=True,size=9,color=RGBColor(0x7D,0xD3,0xFC))
    cr(c[1],v,size=9,color=C_DARK)
pb(doc)

# ── 2. IS PAKETLERI ───────────────────────────────────────────────────────────
hd(doc,"2. Is Paketi Aciklamalari",lvl=1,color=C_BLUE); hline(doc)
for ip in IP:
    t_hdr=doc.add_table(rows=1,cols=3); t_hdr.style='Table Grid'
    c=t_hdr.rows[0].cells; c[0].width=Cm(2.5); c[1].width=Cm(8.5); c[2].width=Cm(10.8)
    for cc in c: shd(cc,ip["renk"]); brd(cc,ip["renk"])
    cr(c[0],ip["kod"],bold=True,size=11,color=C_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
    cr(c[1],ip["ad"],bold=True,size=11,color=C_WHITE)
    m_s=MONTHS[ip["start"]][1]; m_e=MONTHS[ip["end"]][1]
    ms_str=", ".join(MONTHS[m][0] for m in ip["ms"])
    cr(c[2],f"{m_s} - {m_e}  |  Milestonlar: {ms_str}",size=8,color=RGBColor(0xBE,0xDB,0xF7),align=WD_ALIGN_PARAGRAPH.RIGHT)
    t_body=doc.add_table(rows=1,cols=3); t_body.style='Table Grid'
    cc=t_body.rows[0].cells; cc[0].width=Cm(7.5); cc[1].width=Cm(7.5); cc[2].width=Cm(6.8)
    shd(cc[0],"F8FAFC"); shd(cc[1],"F8FAFC"); shd(cc[2],"F0FDF4")
    brd(cc[0],"E2E8F0"); brd(cc[1],"E2E8F0"); brd(cc[2],"BBF7D0")
    p0=cc[0].paragraphs[0]; rh=p0.add_run("Amac & Kapsam\n"); rh.bold=True; rh.font.size=Pt(8); rh.font.name='Calibri'; rh.font.color.rgb=C_BLUE; rb=p0.add_run(ip["ozet"]); rb.font.size=Pt(8); rb.font.name='Calibri'; rb.font.color.rgb=C_DARK
    p1=cc[1].paragraphs[0]; rh2=p1.add_run("Ciktilar\n"); rh2.bold=True; rh2.font.size=Pt(8); rh2.font.name='Calibri'; rh2.font.color.rgb=C_BLUE
    for ckt in ip["ciktilar"]: rb2=p1.add_run(f"  + {ckt}\n"); rb2.font.size=Pt(7.5); rb2.font.name='Calibri'; rb2.font.color.rgb=C_DARK
    p2=cc[2].paragraphs[0]; rh3=p2.add_run("Kabul Kriteri\n"); rh3.bold=True; rh3.font.size=Pt(8); rh3.font.name='Calibri'; rh3.font.color.rgb=C_GREEN; rb3=p2.add_run(ip["kta"]); rb3.font.size=Pt(8); rb3.font.name='Calibri'; rb3.font.color.rgb=C_DARK
    doc.add_paragraph().paragraph_format.space_after=Pt(6)
pb(doc)

# ── 3. GANTT TABLOSU ──────────────────────────────────────────────────────────
hd(doc,"3. 18 Aylik Gantt Tablosu",lvl=1,color=C_BLUE); hline(doc)
para(doc,"Mavi = aktif gelistirme  |  M (kirmizi) = Milestone  |  Bos = beklemede",size=8,color=C_GRAY,sa=6)
COL_IP=Cm(3.8); COL_MON=Cm(1.23)
n_rows=2+len(IP); t_gantt=doc.add_table(rows=n_rows,cols=19); t_gantt.style='Table Grid'
for r in range(n_rows):
    cells=t_gantt.rows[r].cells; cells[0].width=COL_IP
    for col in range(1,19): cells[col].width=COL_MON

cells0=t_gantt.rows[0].cells
shd(cells0[0],HEADER_BG); brd(cells0[0],"334155"); cr(cells0[0],"Is Paketi",bold=True,size=7,color=RGBColor(0x7D,0xD3,0xFC))
for col in range(1,19):
    c=cells0[col]; shd(c,HEADER_BG); brd(c,"334155"); cr(c,MONTHS[col-1][0],bold=True,size=7,color=RGBColor(0x7D,0xD3,0xFC),align=WD_ALIGN_PARAGRAPH.CENTER)

cells1=t_gantt.rows[1].cells
shd(cells1[0],SUBHDR_BG); brd(cells1[0],"334155"); cr(cells1[0],"Donem",bold=True,size=6.5,color=RGBColor(0x94,0xA3,0xB8))
for col in range(1,19):
    c=cells1[col]; shd(c,SUBHDR_BG); brd(c,"334155"); cr(c,MONTHS[col-1][1],size=6.5,color=RGBColor(0x94,0xA3,0xB8),align=WD_ALIGN_PARAGRAPH.CENTER)

for row_i,ip in enumerate(IP):
    actual_row=row_i+2; row_cells=t_gantt.rows[actual_row].cells
    bg_ip="F0F9FF" if row_i%2==0 else "FFFFFF"
    shd(row_cells[0],"1E293B"); brd(row_cells[0],"2E4057")
    p=row_cells[0].paragraphs[0]; r1=p.add_run(f"{ip['kod']}\n"); r1.bold=True; r1.font.size=Pt(7); r1.font.name='Calibri'; r1.font.color.rgb=RGBColor(0x7D,0xD3,0xFC); r2=p.add_run(ip['ad']); r2.font.size=Pt(6.5); r2.font.name='Calibri'; r2.font.color.rgb=RGBColor(0x94,0xA3,0xB8)
    for col in range(1,19):
        m_idx=col-1; c=row_cells[col]; is_active=ip["start"]<=m_idx<=ip["end"]; is_ms=m_idx in ip["ms"]
        if is_ms: shd(c,GANTT_MS); brd(c,GANTT_MS); cr(c,"M",bold=True,size=8,color=C_WHITE,align=WD_ALIGN_PARAGRAPH.CENTER)
        elif is_active: shd(c,GANTT_BLUE); brd(c,GANTT_BLUE)
        else: shd(c,bg_ip); brd(c,"E2E8F0")
pb(doc)

# ── 4. AYLIK AKTIVITE PLANI ───────────────────────────────────────────────────
hd(doc,"4. Aylik Aktivite Plani",lvl=1,color=C_BLUE); hline(doc)
for m_idx in range(18):
    ay_num,ay_ad=MONTHS[m_idx]; aktif_ipler=[ip for ip in IP if ip["start"]<=m_idx<=ip["end"]]
    if not aktif_ipler: continue
    t_ay=doc.add_table(rows=1,cols=2); t_ay.style='Table Grid'
    c=t_ay.rows[0].cells; c[0].width=Cm(3.0); c[1].width=Cm(23.8)
    shd(c[0],"1E293B"); shd(c[1],"1E293B"); brd(c[0],"334155"); brd(c[1],"334155")
    cr(c[0],ay_num,bold=True,size=11,color=C_BLUE,align=WD_ALIGN_PARAGRAPH.CENTER)
    aktif_str="  |  ".join(f"{ip['kod']}" for ip in aktif_ipler)
    ms_bu_ay=[ip for ip in IP if m_idx in ip["ms"]]
    ms_str="  *** MILESTONE: "+" | ".join(f"{ip['kod']}" for ip in ms_bu_ay) if ms_bu_ay else ""
    cr(c[1],f"{ay_ad}   Aktif: {aktif_str}{ms_str}",bold=True,size=9,color=RGBColor(0x7D,0xD3,0xFC))
    n_aktif=len(aktif_ipler); col_w=Cm(26.8/n_aktif)
    t_act=doc.add_table(rows=1,cols=n_aktif); t_act.style='Table Grid'
    for col_i,ip in enumerate(aktif_ipler):
        c=t_act.rows[0].cells[col_i]; c.width=col_w; is_ms=m_idx in ip["ms"]
        bg="FFF5F5" if is_ms else ("F0F9FF" if col_i%2==0 else "F8FAFC")
        shd(c,bg); brd(c,"E2E8F0")
        p=c.paragraphs[0]; rh=p.add_run(f"[{ip['kod']}] {ip['ad']}\n"); rh.bold=True; rh.font.size=Pt(7.5); rh.font.name='Calibri'; rh.font.color.rgb=RGBColor(0xDC,0x26,0x26) if is_ms else C_BLUE
        tasks=ip["aylik"].get(m_idx,["(Bu ay: devam eden gelistirme)"])
        for task in tasks:
            marker=">>> " if is_ms else "  + "; rb=p.add_run(f"{marker}{task}\n"); rb.font.size=Pt(7); rb.font.name='Calibri'; rb.font.color.rgb=C_RED if is_ms else C_DARK
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
pb(doc)

# ── 5. MILESTONE OZETI ────────────────────────────────────────────────────────
hd(doc,"5. Milestone Ozeti",lvl=1,color=C_BLUE); hline(doc)
milestones=[]
for ip in IP:
    for m_idx in ip["ms"]:
        tasks=ip["aylik"].get(m_idx,[])
        ms_task=next((t for t in tasks),"(Milestone tamamlandi)")
        milestones.append({"ms_no":f"M{len(milestones)+1}","ay":MONTHS[m_idx][0],"donem":MONTHS[m_idx][1],"ip":ip["kod"],"ad":ip["ad"],"aciklama":ms_task})
ms_header=["MS#","Ay","Donem","IP","Is Paketi","Milestone Tanimi"]
ms_widths=[Cm(1.2),Cm(1.5),Cm(2.0),Cm(1.5),Cm(4.5),Cm(16.1)]
t_ms=doc.add_table(rows=len(milestones)+1,cols=6); t_ms.style='Table Grid'
for j,(h,w) in enumerate(zip(ms_header,ms_widths)):
    c=t_ms.rows[0].cells[j]; c.width=w; shd(c,HEADER_BG); brd(c,"334155"); cr(c,h,bold=True,size=8,color=RGBColor(0x7D,0xD3,0xFC),align=WD_ALIGN_PARAGRAPH.CENTER)
for i,ms in enumerate(milestones):
    row=t_ms.rows[i+1].cells; vals=[ms["ms_no"],ms["ay"],ms["donem"],ms["ip"],ms["ad"],ms["aciklama"]]
    bg="FFF5F5" if i%2==0 else "FFFFFF"
    for j,(v,w) in enumerate(zip(vals,ms_widths)):
        row[j].width=w; shd(row[j],bg); brd(row[j],"FED7D7")
        aln=WD_ALIGN_PARAGRAPH.CENTER if j<4 else WD_ALIGN_PARAGRAPH.LEFT
        clr=RGBColor(0xDC,0x26,0x26) if j==0 else C_DARK
        cr(row[j],v,bold=(j==0),size=8,color=clr,align=aln)
pb(doc)

# ── 6. RISK MATRISI ───────────────────────────────────────────────────────────
hd(doc,"6. Risk & Bagimlilik Matrisi",lvl=1,color=C_BLUE); hline(doc)
risks=[
    ("Alan","Risk Tanimi","Olasilik","Etki","Seviye","Azaltma Plani"),
    ("Teknik","metrics.py router baglantisi eksik: bazi endpoint'ler erisilemez","Yuksek","Orta","ORTA","Sprint 4'te router.py'ye include_router eklenmesi zorunlu"),
    ("Teknik","Log.tsx App.tsx routing'inden disarda: navigasyonda yok","Yuksek","Dusuk","DUSUK","IP3 icerisinde route + Layout nav girisi eklenir"),
    ("Mimari","Celery var ama asyncio tercih edildi: uzun gorevler bloklar","Orta","Yuksek","ORTA","IP7+ icin run_in_executor zorunlu; IP8'de Celery migrasyonu degerlendirilir"),
    ("Guvenlik","SECRET_KEY eksikligi sureci sonlandiriyor","Orta","Yuksek","YUKSEK","Deployment guide + zorunlu .env sablonu; CI'da env kontrol adimi"),
    ("Bagimlilik","Ollama yerel kurulum: CI/CD ortaminda AI testleri calismayabilir","Orta","Orta","ORTA","Multi-LLM fallback + Ollama mock sinifi ile CI testi"),
    ("Kapsam","VMware/oVirt API erisim gecikmesi IP2'yi erteleyebilir","Orta","Orta","ORTA","Mock client ile paralel gelistirme; gercek API IP2 Sprint 2'ye ertelenir"),
    ("Kaynak","Tek gelistirici varsayimiyla IP7-8 kayabilir","Dusuk","Yuksek","ORTA","IP5 bitiminde kaynak revizyonu; kritik olmayan ozellikler ertelenebilir"),
]
r_w=[Cm(2.2),Cm(7.5),Cm(2),Cm(2),Cm(2),Cm(7.1)]
t_risk=doc.add_table(rows=len(risks),cols=6); t_risk.style='Table Grid'
for i,row in enumerate(risks):
    cells=t_risk.rows[i].cells
    for j,w in enumerate(r_w): cells[j].width=w
    if i==0:
        for j in range(6): shd(cells[j],HEADER_BG); brd(cells[j],"334155"); cr(cells[j],row[j],bold=True,size=8,color=RGBColor(0x7D,0xD3,0xFC),align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        bg="FFFBEB" if i%2==0 else "FFFFFF"
        for j in range(6): shd(cells[j],bg); brd(cells[j],"FDE68A")
        for j,v in enumerate(row):
            clr=(RGBColor(0xDC,0x26,0x26) if "YUKSEK" in v else RGBColor(0xEA,0x58,0x0C) if "ORTA" in v else C_GREEN if "DUSUK" in v else C_DARK) if j==4 else C_DARK
            cr(cells[j],v,bold=(j==4),size=8,color=clr)
pb(doc)

# ── 7. KAYNAK & EFOR ──────────────────────────────────────────────────────────
hd(doc,"7. Kaynak & Efor Ozeti",lvl=1,color=C_BLUE); hline(doc)
para(doc,"1 gun = 8 is saati  |  Tek gelistirici senaryosu  |  +/- %%20 sapma olasidir.",size=8,color=C_GRAY,sa=8)
efor=[
    ("IP","Is Paketi","Donem","Sure","Tahmini Efor","Sprint"),
    ("IP1","Temel Platform Altyapisi","Sub-Mar 2026","2 ay","~37 gun","4"),
    ("IP2","Envanter ve Kesif","Mar-May 2026","3 ay","~30 gun","6"),
    ("IP3","Log Toplama & Analiz","Nis-Haz 2026","3 ay","~25 gun","6"),
    ("IP4","Anomali Tespiti","May-Eyl 2026","5 ay","~40 gun","10"),
    ("IP5","Kapasite Planlama","Tem-Eki 2026","4 ay","~30 gun","8"),
    ("IP6","Servis Haritalama","Eyl 2026-Oca 2027","5 ay","~35 gun","10"),
    ("IP7","AI Chat & Otonom Operasyon","Ara 2026-May 2027","6 ay","~50 gun","12"),
    ("IP8","Test & Urunlestirme","Mar-Tem 2027","5 ay","~50 gun","10"),
    ("TOPLAM","--","Sub 2026-Tem 2027","18 ay","~297 gun","66"),
]
e_w=[Cm(1.8),Cm(5.5),Cm(4.0),Cm(2.0),Cm(3.0),Cm(2.5)]
t_efor=doc.add_table(rows=len(efor),cols=6); t_efor.style='Table Grid'
for i,row in enumerate(efor):
    cells=t_efor.rows[i].cells
    for j,w in enumerate(e_w): cells[j].width=w
    if i==0:
        for j in range(6): shd(cells[j],HEADER_BG); brd(cells[j],"334155"); cr(cells[j],row[j],bold=True,size=8,color=RGBColor(0x7D,0xD3,0xFC),align=WD_ALIGN_PARAGRAPH.CENTER)
    elif i==len(efor)-1:
        for j in range(6): shd(cells[j],"1E293B"); brd(cells[j],"334155"); cr(cells[j],row[j],bold=True,size=9,color=RGBColor(0x34,0xD3,0x99),align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        bg="F8FAFC" if i%2==0 else "FFFFFF"
        for j in range(6): shd(cells[j],bg); brd(cells[j],"E2E8F0"); cr(cells[j],row[j],size=8.5,color=C_DARK)
hline(doc)
para(doc,"Not: Efor tahminleri mevcut kod tabanina dayali analitik ongoru olup +/- %%20 sapma icermektedir. Her is paketi basinda backlog revizyonu onerilen bir uygulamadir.",size=8,italic=True,color=C_GRAY,sa=6)
hline(doc)
para(doc,"datatem AI  |  18 Aylik Proje Plani  |  v1.0  |  Mart 2026",size=8,italic=True,color=C_GRAY,align=WD_ALIGN_PARAGRAPH.CENTER)

output="/home/datatem/ainew/datatem_AI_18ay_proje_plani.docx"
doc.save(output)
print(f"OK: {output}")
